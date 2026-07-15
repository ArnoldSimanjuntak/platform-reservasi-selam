from __future__ import annotations

import json
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


SOURCE = Path(r"D:\Skripsi\Laporan Proposal Arnold Simanjuntak(backup).docx")
OUT = Path(r"C:\Users\arnol\skripsialan\.tmp\laporan-audit")
MEDIA = OUT / "media"


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    MEDIA.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)

    # Extract all embedded media without modifying the source document.
    with zipfile.ZipFile(SOURCE) as archive:
        for name in archive.namelist():
            if name.startswith("word/media/") and not name.endswith("/"):
                (MEDIA / Path(name).name).write_bytes(archive.read(name))

    # Export the complete Chapter IV-V paragraph stream plus table contents.
    lines: list[str] = []
    in_scope = False
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "HASIL DAN PEMBAHASAN":
            in_scope = True
        if in_scope:
            style = paragraph.style.name if paragraph.style else ""
            lines.append(f"P{idx:04d} [{style}] {text}")
        if text == "DAFTAR PUSTAKA":
            break

    lines.append("\n=== TABLES ===")
    for t_idx, table in enumerate(doc.tables, start=1):
        rows = [[cell.text.replace("\n", " / ").strip() for cell in row.cells] for row in table.rows]
        joined = " ".join(" ".join(row) for row in rows).upper()
        if any(term in joined for term in ("BOOKING", "PROVIDER", "PWA", "BLACK BOX", "USERS", "SERVICES", "DIVE_SITES", "RESOURCE")):
            lines.append(f"\nTABLE {t_idx}")
            for row in rows:
                lines.append(" | ".join(row))
    (OUT / "bab4-5-structure.txt").write_text("\n".join(lines), encoding="utf-8")

    # Map image-containing paragraphs to relationship targets.
    image_map = []
    for p_idx, paragraph in enumerate(doc.paragraphs):
        for drawing in paragraph._p.xpath(".//a:blip"):
            rel_id = drawing.get(qn("r:embed"))
            rel = doc.part.rels.get(rel_id)
            image_map.append(
                {
                    "paragraph": p_idx,
                    "paragraph_text": paragraph.text.strip(),
                    "relationship": rel_id,
                    "target": Path(rel.target_ref).name if rel else None,
                }
            )
    (OUT / "image-map.json").write_text(json.dumps(image_map, indent=2), encoding="utf-8")

    # Create a labeled contact sheet for fast visual inspection of embedded media.
    files = sorted(MEDIA.iterdir())
    thumbs = []
    for path in files:
        try:
            im = Image.open(path).convert("RGB")
            original_size = im.size
            im.thumbnail((300, 190))
            thumbs.append((path.name, original_size, im.copy()))
        except Exception:
            continue
    cols, cell_w, cell_h = 3, 340, 245
    rows = (len(thumbs) + cols - 1) // cols
    sheet = Image.new("RGB", (cols * cell_w, rows * cell_h), "white")
    draw = ImageDraw.Draw(sheet)
    font = ImageFont.load_default()
    for i, (name, original_size, im) in enumerate(thumbs):
        x = (i % cols) * cell_w + 20
        y = (i // cols) * cell_h + 35
        sheet.paste(im, (x, y))
        draw.text((x, 8 + (i // cols) * cell_h), f"{name} {original_size[0]}x{original_size[1]}", fill="black", font=font)
    sheet.save(OUT / "media-contact-sheet.png")

    print(f"scope_text={OUT / 'bab4-5-structure.txt'}")
    print(f"media_count={len(files)}")
    print(f"image_map_count={len(image_map)}")
    print(f"contact_sheet={OUT / 'media-contact-sheet.png'}")


if __name__ == "__main__":
    main()
