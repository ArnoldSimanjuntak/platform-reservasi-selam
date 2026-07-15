from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path(r"D:\Skripsi\Laporan Proposal Arnold Simanjuntak(backup).docx")
OUTPUT = Path(r"C:\Users\arnol\skripsialan\docs\Laporan Proposal Arnold Simanjuntak - Revisi Bab V.docx")


KESIMPULAN_INTRO = (
    "Kesimpulan penelitian ini disusun untuk menjawab rumusan masalah berdasarkan hasil "
    "perancangan, implementasi, dan pengujian prototipe SulutDive. Berdasarkan hasil penelitian "
    "yang telah diuraikan pada bab sebelumnya, diperoleh kesimpulan sebagai berikut:"
)

KESIMPULAN_ITEMS = [
    (
        "Rumusan masalah pertama mengenai perancangan aplikasi agregator booking wisata selam "
        "berbasis Progressive Web App di kawasan Lembeh telah dijawab melalui analisis kebutuhan "
        "tiga aktor, pemodelan proses, perancangan basis data, dan perancangan antarmuka. "
        "Perancangan tersebut menghasilkan model alur terpusat yang menghubungkan informasi "
        "layanan, proses booking, serta pengelolaan aplikasi dengan pembagian tanggung jawab yang "
        "jelas antara customer, provider, dan admin."
    ),
    (
        "Rumusan masalah kedua mengenai implementasi fitur utama telah dijawab melalui penerapan "
        "proses yang saling terintegrasi sesuai hak akses setiap aktor. Customer dapat mencari "
        "layanan dan melakukan booking, provider yang telah diverifikasi dapat mengelola layanan "
        "serta pesanan, sedangkan admin dapat memverifikasi provider dan memantau data aplikasi. "
        "Dengan demikian, kebutuhan fungsional utama ketiga aktor telah diterapkan dalam satu "
        "alur sistem sesuai dengan ruang lingkup penelitian."
    ),
    (
        "Rumusan masalah ketiga mengenai penerapan Progressive Web App telah dijawab melalui "
        "penggunaan web app manifest, service worker, tampilan responsif, kemampuan instalasi pada "
        "perangkat mobile, dan halaman fallback ketika koneksi tidak tersedia. Penerapan tersebut "
        "memungkinkan SulutDive diakses melalui browser maupun dipasang pada perangkat mobile. "
        "Namun, dukungan offline dibatasi pada halaman fallback dan aset yang telah disimpan; proses "
        "transaksional seperti booking, unggah bukti pembayaran, dan perubahan data tetap "
        "memerlukan koneksi internet."
    ),
    (
        "Berdasarkan pengujian black box dan pengujian karakteristik PWA yang dilakukan, fungsi "
        "utama dapat dijalankan sesuai dengan skenario pengujian dan komponen utama PWA dapat "
        "berfungsi pada ruang lingkup yang ditetapkan. Hasil tersebut mendukung bahwa prototipe "
        "telah memenuhi tujuan penelitian pada tingkat fungsional, tetapi belum dapat diartikan "
        "sebagai bukti mengenai tingkat penerimaan pengguna, dampak operasional, atau kesiapan "
        "aplikasi untuk digunakan dalam lingkungan operasional yang sebenarnya."
    ),
]

KESIMPULAN_CLOSING = (
    "Secara keseluruhan, penelitian ini menghasilkan prototipe agregator booking wisata selam "
    "berbasis PWA yang dapat menjadi model awal integrasi layanan wisata selam di kawasan Lembeh. "
    "Kesimpulan tersebut berlaku sesuai batasan penelitian dan hasil pengujian yang telah dilakukan."
)

SARAN_INTRO = (
    "Berdasarkan hasil penelitian dan keterbatasan yang masih terdapat pada penelitian ini, "
    "beberapa saran untuk pengembangan dan penelitian selanjutnya adalah sebagai berikut:"
)

SARAN_LEAD = (
    "Saran berikut diarahkan untuk memperkuat validitas hasil penelitian, kesesuaian sistem dengan "
    "kondisi lapangan, serta kesiapan prototipe sebelum diterapkan secara lebih luas."
)

SARAN_ITEMS = [
    (
        "Penelitian selanjutnya perlu melibatkan customer, provider, dan admin yang representatif "
        "dalam pengujian pengguna. Evaluasi dapat mengukur keberhasilan penyelesaian tugas, waktu "
        "penyelesaian, jumlah kesalahan, dan tingkat usability, misalnya menggunakan System "
        "Usability Scale, agar hasil penelitian tidak hanya didasarkan pada keberhasilan fungsi."
    ),
    (
        "Pengujian PWA perlu diperluas pada beberapa perangkat, sistem operasi, browser, ukuran "
        "layar, dan kondisi jaringan. Setiap pengujian sebaiknya dilengkapi dengan informasi "
        "perangkat, versi browser, skenario jaringan, dan bukti hasil agar kemampuan instalasi, "
        "responsivitas, serta perilaku offline dapat dievaluasi secara lebih terukur."
    ),
    (
        "Model kapasitas layanan dan ketersediaan peralatan perlu divalidasi menggunakan data "
        "operasional provider di kawasan Lembeh. Kapasitas operasional layanan juga perlu dibedakan "
        "secara tegas dari daya dukung ekologis lokasi penyelaman; penelitian lanjutan dapat "
        "mengintegrasikan data daya dukung lingkungan apabila data dan standar dari pihak terkait "
        "telah tersedia."
    ),
    (
        "Kriteria verifikasi provider perlu dikaji bersama instansi pariwisata, asosiasi usaha, "
        "atau pihak berwenang agar jenis dokumen, masa berlaku, proses pembaruan, dan alasan "
        "penolakan memiliki dasar yang sesuai dengan praktik di lapangan."
    ),
    (
        "Sebelum diterapkan secara nyata, perlu dilakukan evaluasi keamanan dan privasi terhadap "
        "data akun, dokumen provider, dan bukti pembayaran. Evaluasi dapat mencakup pengujian hak "
        "akses berbasis peran, kebijakan akses basis data, penyimpanan berkas privat, validasi "
        "unggahan, pencatatan aktivitas, serta prosedur pencadangan dan pemulihan data."
    ),
    (
        "Pengembangan seperti integrasi payment gateway, notifikasi otomatis, ulasan layanan, dan "
        "analisis data dapat dilakukan setelah kebutuhan pengguna dan proses bisnis inti divalidasi. "
        "Penelitian berikutnya sebaiknya tidak hanya menilai apakah fitur tersebut berhasil dibuat, "
        "tetapi juga mengukur pengaruhnya terhadap kemudahan proses booking, kecepatan layanan, dan "
        "kepercayaan pengguna."
    ),
]

SARAN_CLOSING = (
    "Dengan melaksanakan saran tersebut, penelitian lanjutan diharapkan dapat menghasilkan sistem "
    "yang tidak hanya berfungsi secara teknis, tetapi juga tervalidasi dari sisi pengguna, proses "
    "operasional, keamanan, dan konteks pengelolaan wisata selam di kawasan Lembeh."
)


def replace_paragraph_text(paragraph, text):
    """Replace text while preserving paragraph properties and representative run formatting."""
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[0]._r.rPr)

    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)

    if text:
        run = paragraph.add_run(text)
        if run_properties is not None:
            run._r.insert(0, run_properties)


def delete_paragraph(paragraph):
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)
    paragraph._p = paragraph._element = None


def find_last_heading(paragraphs, text):
    matches = [i for i, p in enumerate(paragraphs) if p.text.strip().casefold() == text.casefold()]
    if not matches:
        raise ValueError(f"Heading tidak ditemukan: {text}")
    return matches[-1]


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)
    paragraphs = doc.paragraphs

    kesimpulan_idx = find_last_heading(paragraphs, "Kesimpulan")
    saran_idx = find_last_heading(paragraphs, "Saran")
    daftar_idx = find_last_heading(paragraphs, "DAFTAR PUSTAKA")
    if not (kesimpulan_idx < saran_idx < daftar_idx):
        raise ValueError("Urutan bagian Bab V tidak sesuai dengan struktur yang diharapkan.")

    conclusion_region = paragraphs[kesimpulan_idx + 1:saran_idx]
    if len(conclusion_region) < 7:
        raise ValueError("Paragraf bagian Kesimpulan tidak mencukupi untuk revisi terarah.")

    replace_paragraph_text(conclusion_region[0], KESIMPULAN_INTRO)
    conclusion_region[0].style = doc.styles["Normal"]

    for paragraph, text in zip(conclusion_region[1:5], KESIMPULAN_ITEMS):
        replace_paragraph_text(paragraph, text)
        paragraph.style = doc.styles["List Paragraph"]

    replace_paragraph_text(conclusion_region[-1], KESIMPULAN_CLOSING)
    conclusion_region[-1].style = doc.styles["Normal"]

    for paragraph in conclusion_region[5:-1]:
        delete_paragraph(paragraph)

    # Refresh paragraph collection because two old conclusion items were removed.
    paragraphs = doc.paragraphs
    saran_idx = find_last_heading(paragraphs, "Saran")
    daftar_idx = find_last_heading(paragraphs, "DAFTAR PUSTAKA")
    suggestion_region = paragraphs[saran_idx + 1:daftar_idx]
    nonempty_suggestion_region = [p for p in suggestion_region if p.text.strip()]
    if len(nonempty_suggestion_region) < 9:
        raise ValueError("Paragraf bagian Saran tidak mencukupi untuk revisi terarah.")

    replace_paragraph_text(nonempty_suggestion_region[0], SARAN_INTRO)
    nonempty_suggestion_region[0].style = doc.styles["Normal"]
    replace_paragraph_text(nonempty_suggestion_region[1], SARAN_LEAD)
    nonempty_suggestion_region[1].style = doc.styles["Normal"]

    for paragraph, text in zip(nonempty_suggestion_region[2:8], SARAN_ITEMS):
        replace_paragraph_text(paragraph, text)
        paragraph.style = doc.styles["List Paragraph"]

    replace_paragraph_text(nonempty_suggestion_region[8], SARAN_CLOSING)
    nonempty_suggestion_region[8].style = doc.styles["Normal"]

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
